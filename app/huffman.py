####################################################################
# Implementation of the Huffman Coding Algorithm.                  #
# Parts of this implementation were borrowed from other sources.   #
# See specific comments for who authored what.                     #
#                                                                  #
# Written by Lewis Kim.                                            #
####################################################################

from itertools import groupby
from heapq import *
import os
from datetime import datetime

from bitstring import BitArray, BitStream
from bitarray import bitarray

from app.utils import *
from app.AESCipher import AESCipher
from docx import *

# Node class that makes up a binary tree. Used to constuct a Huffman Tree.
class Node:
	left = None
	right = None
	item = None
	value = 0

	def __init__(self, item, value):
		self.item = item
		self.value = value

	def set_children(self, left_node, right_node):
		self.left = left_node
		self.right = right_node

	def __lt__(self, other):
		return self.value < other.value


# Huffman Coder that encodes file content, and compresses/decompresses it.
# Designed to be used around FileCase's needs, e.g. checking directories in FileCase/data.
class HuffmanCoder:

	"""Create a new instance of HuffmanCoder with a FILEPATH of a text file.

	   class variables:
	   - self.filepath: Filepath of the file to be encoded/compress.
	   - self.destination_path: 
	   - self.filename: The name of the file from self.filepath, without the file extension.
	   - self.file_ext: The file extension of the file from self.filepath.
	   - self.bit_length: Length of the encoded file content (i.e. length of bit string).
	   - self.content: Text content of the selected file.
	   - self.char_to_code: Dictionary of the Huffman Codes. Keys are characters appearing
	     in the file, and the values are the codes associated with those characters.
	   - self.code_to_char: key:value reversal of self.char_to_code.
	   - self.freq_tree: Dictionary whose keys are characters of the content of the file.
	     and values the frequencies of those characters.
	   - self.encrypted: Whether or not the content of the file is to be encrypted."""
	def __init__(self, filepath, destination_path, encrypted=False):
		self.filepath = filepath
		self.destination_path = destination_path

		self.filename = self.filepath.split("/")[-1].split(".")[0]
		self.file_ext = os.path.splitext(filepath)[1]
		self.bit_length = 0

		self.content = ""

		self.char_to_code = {}
		self.code_to_char = {}

		if self.file_ext == ".txt":
			with open(self.filepath, 'r+') as file:
				self.content = file.read()

		elif self.file_ext == ".docx":
			document = Document(filepath)

			for para in document.paragraphs:
				self.content += para.text

		# May not implement; may require separate class.
		elif self.file_ext == ".jpg" or self.file_ext == ".jpeg":
			return

		else:
			raise ValueError("Wrong File Type for Huffman Coding.")
			return

		self.freq_tree = self.build_freq_tree(self.content)


	"""Build a character frequency dictionary, whose keys are the characters of the string
	   CONTENT, and values are the frequency of those characters in CONTENT."""
	def build_freq_tree(self, content):
		freq = {}
		#
		for char in content:
			if not char in freq:
				freq[char] = 0

			freq[char] += 1

		return freq


	"""Encode the characters of TO_ENCODE using the Huffman Coding Algorithm.
	   This function constructs a Huffman Tree, and returns a dictionary whose keys are characters
	   in self.content, and values are the bit code associated with that character.

	   Help From: https://www.techrepublic.com/article/huffman-coding-in-python/"""
	def encode(self):
		codes = {}

		def codeIt(s, node):
			if node.item:
				if not s:
					codes[node.item] = "0"
				else:
					codes[node.item] = s
			else:
				codeIt(s+"0", node.left)
				codeIt(s+"1", node.right)

		queue = [Node(a, len(list(b))) for a, b in groupby(sorted(self.content))]
		heapify(queue)

		while len(queue) > 1:
			left = heappop(queue)
			right = heappop(queue)

			node = Node(None, right.value + left.value)
			node.set_children(left, right)
			heappush(queue, node)

		codeIt("", queue[0])

		return codes


	"""Encode the contents of the selected file (pointer: self.content)."""
	def compress(self):
		compressed_filepath = self.destination_path + "/" + self.filename + "-" + self.file_ext[1:] + ".bin"

		# Coding for .txt and .docx files.
		if self.file_ext == ".txt" or ".docx":

			self.char_to_code = self.encode()
			self.code_to_char = {value:key for key, value in self.char_to_code.items()}
			encoded_content = "".join([self.char_to_code[a] for a in self.content])

			self.bit_length = len(encoded_content)
			bit_array = BitArray(bin = encoded_content)

			with open(compressed_filepath, 'wb') as compressed_file:
				compressed_file.write(bit_array.tobytes())
				compressed_file.close()

			return compressed_filepath

		else:
			raise ValueError("HuffmanCoder Compression Error: Wrong File Type")
			return


	"""Decompress the .bin file located at COMPRESSED_FILEPATH to .txt or .docx, depending on the 
	   original file (determined by self.file_ext when the instance of this HuffmanCoder was created).
	   New decompressed file written to DECOMPRESSION_PATH.
	   Returns the name of the file at DECOMPRESSION_PATH with the correct extention."""
	def decompress(self, compressed_filepath, decompression_path):
		bits = BitArray(filename = compressed_filepath)

		encoded_content = bits.bin[0:self.bit_length]
		decoder = {k:bitarray(v) for k, v in self.char_to_code.items()}

		decoded_chars = bitarray(encoded_content).decode(decoder)
		decoded_content = ''.join(x for x in decoded_chars)

		if self.file_ext == ".txt":
			decompressed_filepath = decompression_path + "/" + self.filename + "_(decompressed)" + ".txt"

			with open(decompressed_filepath, 'w+') as decompressed_file:
				decompressed_file.write(decoded_content)
				decompressed_file.close()

			return decompressed_filepath

		if self.file_ext == ".docx":
			decompressed_filepath = decompression_path + "/" + self.filename + "_(decompressed)" + ".docx"

			document = Document()
			document.add_paragraph(decoded_content)
			document.save(decompressed_filepath)

			return decompressed_filepath


	"""Print some attributes of this instance of HuffmanEncoder. Used for testing purposes to see
	   if attrs point properly to the correct data."""
	def print_attrs(self):
		print("Filepath: " + self.filepath)
		print("File Name: " + self.filename)
		print("File Extension: " + self.file_ext)
		print(self.freq_tree)

